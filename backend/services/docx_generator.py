"""Word document report generator."""

import base64
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from models.cluster import ClusterResult
from models.report import MetricsResult, ReportResult


class DocxGenerator:
    """Service for generating Word document reports."""

    def __init__(self) -> None:
        """Initialize DOCX generator."""
        pass

    def generate(
        self,
        report: ReportResult,
        clusters: list[ClusterResult],
        chart_pngs: dict[str, str],
        output_path: str,
    ) -> str:
        """
        Generate professional DOCX report.

        Args:
            report: Report metadata and executive summary
            clusters: List of cluster results
            chart_pngs: Dict of chart names to base64-encoded PNG data
            output_path: Path to save the DOCX file

        Returns:
            Path to the generated DOCX file
        """
        doc = Document()

        # Title page
        title = doc.add_heading(report.title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Generated: {report.created_at}")
        doc.add_page_break()

        # Executive Summary
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(report.executive_summary)
        doc.add_page_break()

        # Key Metrics table
        doc.add_heading("Key Metrics", level=1)
        metrics = report.metrics

        # Create metrics table
        table = doc.add_table(rows=8, cols=2, style="Light Grid Accent 1")
        table.rows[0].cells[0].text = "Total Incidents"
        table.rows[0].cells[1].text = str(metrics.total_incidents)

        table.rows[1].cells[0].text = "Mean Resolution Time (hours)"
        table.rows[1].cells[1].text = (
            f"{metrics.mean_resolution_hours:.1f}"
            if metrics.mean_resolution_hours
            else "N/A"
        )

        table.rows[2].cells[0].text = "Median Resolution Time (hours)"
        table.rows[2].cells[1].text = (
            f"{metrics.median_resolution_hours:.1f}"
            if metrics.median_resolution_hours
            else "N/A"
        )

        table.rows[3].cells[0].text = "P50 Resolution Time (hours)"
        table.rows[3].cells[1].text = (
            f"{metrics.p50_resolution_hours:.1f}"
            if metrics.p50_resolution_hours
            else "N/A"
        )

        table.rows[4].cells[0].text = "P90 Resolution Time (hours)"
        table.rows[4].cells[1].text = (
            f"{metrics.p90_resolution_hours:.1f}"
            if metrics.p90_resolution_hours
            else "N/A"
        )

        table.rows[5].cells[0].text = "SEV1 Incidents"
        table.rows[5].cells[1].text = str(metrics.sev1_count)

        table.rows[6].cells[0].text = "SEV2 Incidents"
        table.rows[6].cells[1].text = str(metrics.sev2_count)

        table.rows[7].cells[0].text = "SEV3/4 Incidents"
        table.rows[7].cells[1].text = str(
            metrics.sev3_count + metrics.sev4_count
        )

        doc.add_page_break()

        # Embed charts as images
        if chart_pngs:
            doc.add_heading("Charts", level=1)
            for chart_name, b64_png in chart_pngs.items():
                # Format chart name nicely
                formatted_name = (
                    chart_name.replace("_", " ")
                    .replace("-", " ")
                    .title()
                )
                doc.add_heading(formatted_name, level=2)

                try:
                    # Decode base64 PNG and embed
                    image_bytes = base64.b64decode(b64_png)
                    doc.add_picture(BytesIO(image_bytes), width=Inches(6.0))
                except Exception as e:
                    import logging
                    logging.error(f"Failed to embed chart '{chart_name}': {e}")
                    # Add visible warning in document
                    warning_para = doc.add_paragraph(f"[Chart image failed to embed: {e}]")
                    warning_para.style = 'Intense Quote'

            doc.add_page_break()

        # Incident Clusters
        if clusters:
            doc.add_heading("Incident Clusters", level=1)
            for cluster in clusters:
                doc.add_heading(
                    f"Cluster: {cluster.summary or f'Cluster {cluster.cluster_id}'}",
                    level=2,
                )
                if cluster.centroid_text:
                    doc.add_paragraph(cluster.centroid_text)
                doc.add_paragraph(f"Incidents: {cluster.size}")

        # Save document
        doc.save(output_path)
        return output_path
